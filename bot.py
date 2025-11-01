import os
import pickle
import telebot
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google_auth_oauthlib.flow import InstalledAppFlow
from docx import Document
from docx.shared import Inches
from datetime import datetime

# === CONFIGURACIÓN ===
BOT_TOKEN = "8341522444:AAE0wQADD1dnL9R38usxFb9wT6h8GSPkDKs"
bot = telebot.TeleBot(BOT_TOKEN)
SCOPES = ["https://www.googleapis.com/auth/drive.file"]
MAIN_FOLDER_ID = "1zii-j7wXNnqJMSTC1HYZ-NX8tVy0al1K"

# === AUTENTICACIÓN GOOGLE DRIVE ===
creds = None
if os.path.exists("token.pkl"):
    with open("token.pkl", "rb") as token:
        creds = pickle.load(token)

if not creds:
    flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
    creds = flow.run_local_server(port=0)
    with open("token.pkl", "wb") as token:
        pickle.dump(creds, token)

drive_service = build("drive", "v3", credentials=creds)

# === DATOS TEMPORALES ===
user_data = {}

# === SUBCARPETAS Y SUS IDS ===
SUBFOLDERS = {
    "LUNES": "1EWpeg2ngKaaudpE2mOyfVdfGStCGVmzf",
    "MARTES": "1LDNkXGCJhceSY6sc7DXHHkDjli6kUkRL",
    "MIERCOLES": "1u-BjthuqKGOgCu0WYrBf7GmQchXS4k7O",
    "JUEVES": "1D3zZLj1lCxrZmnJ56V2YY6BjjydERhQR",
    "VIERNES": "1GgdStu-9RoklJxWfHhBqpLwsZy3KIcBN",
}

# === OBTENER ID DE SUBCARPETA ===
def get_subfolder_id(day_name):
    return SUBFOLDERS.get(day_name.upper())

# === COMANDO /start ===
@bot.message_handler(commands=["start"])
def start(message):
    bot.reply_to(message, "👋 Hola, por favor ingresa el nombre que llevará el informe:")
    bot.register_next_step_handler(message, ask_day_folder)

# === SELECCIONAR CARPETA ===
def ask_day_folder(message):
    user_data[message.chat.id] = {"report_name": message.text}
    markup = telebot.types.ReplyKeyboardMarkup(one_time_keyboard=True, resize_keyboard=True)
    markup.add(*SUBFOLDERS.keys())
    bot.send_message(message.chat.id, "📂 Selecciona la carpeta del día:", reply_markup=markup)
    bot.register_next_step_handler(message, ask_description)

# === DESCRIPCIÓN DE ACTIVIDAD ===
def ask_description(message):
    folder_id = get_subfolder_id(message.text)
    if not folder_id:
        bot.reply_to(message, f"⚠️ No se encontró la carpeta '{message.text}'. Verifica que exista en tu Drive.")
        return
    user_data[message.chat.id]["folder_id"] = folder_id
    bot.send_message(message.chat.id, "✏️ Escribe la descripción del trabajo realizado:")
    bot.register_next_step_handler(message, save_desc)

def save_desc(message):
    user_data[message.chat.id]["desc"] = message.text
    bot.send_message(message.chat.id, "👤 ¿Quién solicitó el trabajo?")
    bot.register_next_step_handler(message, save_solicitante)

def save_solicitante(message):
    user_data[message.chat.id]["solicitante"] = message.text
    bot.send_message(message.chat.id, "🏢 ¿De qué oficina o área proviene la solicitud?")
    bot.register_next_step_handler(message, save_oficina)

def save_oficina(message):
    user_data[message.chat.id]["oficina"] = message.text
    bot.send_message(message.chat.id, "⚙️ Describe el trabajo realizado:")
    bot.register_next_step_handler(message, save_realizado)

def save_realizado(message):
    user_data[message.chat.id]["realizado"] = message.text
    bot.send_message(message.chat.id, "📝 ¿Quién realizó el informe o prestó el servicio?")
    bot.register_next_step_handler(message, save_responsable)

def save_responsable(message):
    user_data[message.chat.id]["responsable"] = message.text
    bot.send_message(message.chat.id, "📝 Observaciones:")
    bot.register_next_step_handler(message, save_observacion)

def save_observacion(message):
    user_data[message.chat.id]["observacion"] = message.text
    bot.send_message(message.chat.id, "📷 Envía la imagen que quieres agregar al informe:")
    bot.register_next_step_handler(message, save_image)

# === GUARDAR IMAGEN ===
def save_image(message):
    if message.content_type != 'photo':
        bot.reply_to(message, "⚠️ Por favor envía una imagen válida.")
        bot.register_next_step_handler(message, save_image)
        return

    # Descargar la imagen
    file_info = bot.get_file(message.photo[-1].file_id)
    downloaded_file = bot.download_file(file_info.file_path)
    image_name = f"{message.chat.id}_imagen.jpg"
    with open(image_name, "wb") as f:
        f.write(downloaded_file)

    user_data[message.chat.id]["image"] = image_name
    create_word_and_upload(message.chat.id, message)

# === CREAR Y SUBIR INFORME ===
def create_word_and_upload(user_id, message):
    data = user_data[user_id]
    fecha = datetime.now().strftime("%d/%m/%Y")
    nombre_informe = data["report_name"] + ".docx"

    doc = Document()
    doc.add_heading("INFORME DE ACTIVIDADES DIARIAS", level=1)
    doc.add_paragraph(f"Fecha del Informe: {fecha}")
    doc.add_paragraph(f"Responsable informe: {message.from_user.first_name}")
    doc.add_paragraph("\nDetalle del Trabajo Realizado\n")
    doc.add_paragraph(f"▲ A. Trabajo 1: {data['desc']}")
    doc.add_paragraph(f"Quien Solicitó: {data['solicitante']}")
    doc.add_paragraph(f"Oficina/Área: {data['oficina']}")
    doc.add_paragraph(f"Trabajo Realizado: {data['realizado']}")
    doc.add_paragraph(f"Quien realizó el informe / prestó el servicio: {data['responsable']}")
    doc.add_paragraph(f"Observación: {data['observacion']}")

    # Agregar imagen si existe
    if "image" in data:
        doc.add_paragraph("\nImagen adjunta:\n")
        doc.add_picture(data["image"], width=Inches(4))

    doc.save(nombre_informe)

    file_metadata = {"name": nombre_informe, "parents": [data["folder_id"]]}
    media = MediaFileUpload(nombre_informe, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    drive_service.files().create(body=file_metadata, media_body=media, fields="id").execute()

    with open(nombre_informe, "rb") as f:
        bot.send_document(message.chat.id, f)

    bot.reply_to(message, "✅ Informe creado y subido correctamente al Drive.")

    # Limpiar archivos temporales
    os.remove(nombre_informe)
    if "image" in data:
        os.remove(data["image"])
    user_data.pop(user_id, None)

# === EJECUCIÓN ===
if __name__ == "__main__":
    print("🤖 Bot ejecutándose...")
    bot.polling()
